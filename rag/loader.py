import os
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_core.documents import Document


def load_pdf(pdf_path: str):
    loader = PyMuPDFLoader(pdf_path)
    docs = loader.load()

    # From your notebook: add page_number = page + 1
    for d in docs:
        if "page" in d.metadata and isinstance(d.metadata["page"], int):
            d.metadata["page_number"] = d.metadata["page"] + 1

    return docs


def load_docx(docx_path: str):
    from docx import Document as DocxDocument

    doc = DocxDocument(docx_path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    return [
        Document(
            page_content=full_text,
            metadata={"source": docx_path, "page_number": 1},
        )
    ]


def load_file(file_path: str):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use .pdf or .docx")
